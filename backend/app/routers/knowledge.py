"""Knowledge base management endpoints."""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pydantic import BaseModel
from typing import Optional
import uuid
import logging

from app.database import get_db
from app.models.db_models import KnowledgeDocument
from app.agents.knowledge_agent import knowledge_agent

router = APIRouter(prefix="/knowledge", tags=["Knowledge Base"])
logger = logging.getLogger(__name__)


class SearchRequest(BaseModel):
    query: str
    equipment_type: Optional[str] = None


class AIAnswerRequest(BaseModel):
    query: str
    equipment_type: Optional[str] = None


@router.post("/upload")
async def upload_document(
    title: str = Form(...),
    document_type: str = Form(...),
    equipment_type: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    content: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db),
):
    """Upload a document to the knowledge base."""
    doc_content = content or ""

    if file:
        file_content = await file.read()
        filename = file.filename or ""
        if filename.endswith(".txt"):
            doc_content = file_content.decode("utf-8", errors="replace")
        elif filename.endswith(".pdf"):
            try:
                import PyPDF2, io
                reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                doc_content = "\n".join(
                    page.extract_text() for page in reader.pages if page.extract_text()
                )
            except Exception as e:
                doc_content = file_content.decode("utf-8", errors="replace")
        elif filename.endswith(".docx"):
            try:
                import docx, io
                doc = docx.Document(io.BytesIO(file_content))
                doc_content = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                doc_content = file_content.decode("utf-8", errors="replace")
        else:
            doc_content = file_content.decode("utf-8", errors="replace")

    if not doc_content.strip():
        raise HTTPException(status_code=400, detail="No content provided")

    doc_id = str(uuid.uuid4())
    db_doc = KnowledgeDocument(
        id=doc_id,
        title=title,
        document_type=document_type,
        equipment_type=equipment_type,
        content=doc_content[:50000],
        file_path=file.filename if file else None,
        doc_metadata={"document_type": document_type, "equipment_type": equipment_type},
    )
    db.add(db_doc)

    chunk_count = await knowledge_agent.add_document(
        title=title,
        content=doc_content,
        metadata={
            "document_type": document_type,
            "equipment_type": equipment_type or "general",
            "doc_id": doc_id,
        },
    )

    db_doc.chunk_count = chunk_count
    db_doc.is_indexed = chunk_count > 0
    await db.commit()

    return {
        "document_id": doc_id,
        "title": title,
        "chunks_indexed": chunk_count,
        "message": f"Document '{title}' uploaded and indexed successfully",
    }


@router.get("/documents")
async def list_documents(db: AsyncSession = Depends(get_db)):
    """List all indexed documents."""
    result = await db.execute(
        select(KnowledgeDocument).order_by(KnowledgeDocument.created_at.desc())
    )
    docs = result.scalars().all()
    return [
        {
            "id": d.id,
            "title": d.title,
            "document_type": d.document_type,
            "equipment_type": d.equipment_type,
            "chunk_count": d.chunk_count,
            "is_indexed": d.is_indexed,
            "created_at": d.created_at.isoformat() if d.created_at else None,
        }
        for d in docs
    ]


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, db: AsyncSession = Depends(get_db)):
    """Delete a document from the knowledge base."""
    result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    await db.delete(doc)
    await db.commit()
    return {"message": f"Document '{doc.title}' deleted successfully"}


@router.post("/search")
async def search_knowledge(request: SearchRequest):
    """Search the knowledge base — accepts JSON body."""
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    results = await knowledge_agent.retrieve(request.query, request.equipment_type)

    return {
        "query": request.query,
        "results": results,
        "count": len(results),
    }


@router.get("/search")
async def search_knowledge_get(
    query: str = Query(..., min_length=1),
    equipment_type: Optional[str] = Query(None),
):
    """Search the knowledge base — accepts GET query param."""
    results = await knowledge_agent.retrieve(query, equipment_type)
    return {
        "query": query,
        "results": results,
        "count": len(results),
    }


@router.post("/ai-answer")
async def get_ai_answer(request: AIAnswerRequest):
    """
    Get a full AI-generated answer grounded in the knowledge base.
    Retrieves relevant context then asks the LLM to answer the question.
    """
    from app.agents.llm_client import get_llm
    from langchain_core.messages import HumanMessage, SystemMessage

    # Retrieve relevant knowledge
    docs = await knowledge_agent.retrieve(request.query, request.equipment_type)
    context = "\n\n".join([f"[{d['source']}]:\n{d['content']}" for d in docs[:3]])

    if not context:
        context = "Use your general knowledge about steel plant maintenance."

    prompt = f"""You are an expert steel plant maintenance engineer AI.

Using the knowledge base context below, answer the question thoroughly and professionally.

FORMATTING RULES:
- Use ## for section headings
- Use numbered lists for sequential steps
- Use bullet points (- item) for non-sequential points
- Use -> for action recommendations
- Do NOT use **double asterisks** for bold
- Do NOT use emojis
- Write in clear, professional engineering language

Knowledge Base Context:
{context}

Question: {request.query}

Structure your answer with:
## Direct Answer
## Key Points
## Relevant Thresholds or Values (if applicable)
## Recommended Actions"""

    try:
        llm = get_llm(temperature=0.1)
        response = await llm.ainvoke([
            SystemMessage(content="You are an expert industrial maintenance AI for Tata Steel."),
            HumanMessage(content=prompt),
        ])
        answer = response.content
    except Exception as e:
        logger.error(f"AI answer failed: {e}")
        answer = f"Knowledge base results for '{request.query}':\n\n" + context[:1000]

    return {
        "query": request.query,
        "answer": answer,
        "sources": [d["source"] for d in docs],
        "context_used": len(docs),
    }
